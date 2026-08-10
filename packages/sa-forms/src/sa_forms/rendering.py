"""Render a completed session into a reviewable document.

Every renderer produces the same parts:

1. **The answers**, grouped by topic.
2. **Agreements** — what was put to the participant, in the words it was put in,
   and what they decided. A submission gathered under terms is only as good as
   the record that they were accepted, and by whom.
3. **Open action items** — including any mandatory field nobody could close.
   The document states what is missing rather than presenting a partial record
   as complete.
4. **Open questions** — anything the participant asked that had to go to a team.
   An approver reading a value that was guessed at needs to know it was, and the
   question is the evidence.
5. **A provenance appendix** — where each answer came from: which channel, which
   person, and the verbatim text that supports it. This is what makes a
   conversationally-gathered document auditable, and it is the part a manually
   filled form never has.

Renderers whose library is absent report themselves unavailable rather than
failing at import, so a deployment can ship Markdown-only if it wants to.
"""

from __future__ import annotations

import io
import json
from abc import ABC, abstractmethod
from datetime import UTC
from typing import Any

from sa_platform.errors import ConfigurationError
from sa_platform.logging import get_logger

from .actions import noted_discrepancies, open_items, recommended_gaps_as_notes
from .coercion import render_value
from .completeness import analyse
from .models import (
    AnswerState,
    ArtifactStatus,
    FormDefinition,
    FormSession,
)

logger = get_logger(__name__)

_REDACTED = "••••••"


def _display(field: Any, answer: Any) -> str:
    """The value as it should appear in a document."""
    if answer is None or answer.state is AnswerState.EMPTY:
        return ""
    if answer.state in (AnswerState.SKIPPED, AnswerState.NOT_APPLICABLE):
        return "(not provided)"
    if field.sensitive:
        return _REDACTED
    return render_value(field, answer.value)


def _stamp(epoch: float) -> str:
    """A decision time an auditor can read, in UTC."""
    from datetime import datetime

    return datetime.fromtimestamp(epoch, tz=UTC).strftime("%Y-%m-%d %H:%M UTC")


def agreement_rows(form: FormDefinition, session: FormSession) -> list[tuple[str, ...]]:
    """Agreement decisions as (title, kind, version, decision, actor, when).

    Built from the session's own records rather than from the definition, so an
    agreement removed or reworded on the form afterwards cannot change what this
    document says was agreed.
    """
    rows: list[tuple[str, ...]] = []
    for record in session.agreements:
        agreement = form.try_agreement(record.agreement_id)
        rows.append(
            (
                agreement.title if agreement else record.agreement_id,
                agreement.kind.value if agreement else "",
                record.version,
                record.decision.value,
                record.actor,
                _stamp(record.decided_at),
            )
        )
    return rows


def open_questions(session: FormSession) -> list[Any]:
    """Questions still sitting with a team."""
    return session.open_support_requests()


def completeness_line(form: FormDefinition, session: FormSession) -> str:
    """How complete this submission is, in the units the form is measured in.

    An agreement form has no fields, so counting them says "100%" about a
    document nobody has agreed to anything on — and an intake form with terms
    attached is not fully described by its field count either.
    """
    report = analyse(form, session)
    parts: list[str] = []
    if report.total_fields:
        parts.append(
            f"{report.mandatory_answered}/{report.mandatory_total} required, "
            f"{report.overall_percent:.0f}% overall"
        )
    if report.agreements_required:
        parts.append(
            f"{report.agreements_accepted}/{report.agreements_required} agreements accepted"
        )
    if report.agreements_declined:
        parts.append(f"{len(report.agreements_declined)} declined")
    return " · ".join(parts) or "nothing to capture"


def _status_note(answer: Any) -> str:
    if answer is None or answer.state is AnswerState.EMPTY:
        return "Not captured"
    return {
        AnswerState.PROPOSED: "Unconfirmed",
        AnswerState.ANSWERED: "Captured",
        AnswerState.CONFIRMED: "Confirmed",
        AnswerState.SKIPPED: "Skipped",
        AnswerState.NOT_APPLICABLE: "Not applicable",
    }.get(answer.state, "")


class Renderer(ABC):
    """Turns a session into bytes."""

    format: str = ""
    media_type: str = "application/octet-stream"
    extension: str = "bin"

    @abstractmethod
    def render(
        self, form: FormDefinition, session: FormSession, *, include_provenance: bool = True
    ) -> bytes:
        ...

    @classmethod
    def available(cls) -> bool:
        """Whether this renderer's dependencies are installed."""
        return True


class JsonRenderer(Renderer):
    """Machine-readable export. Always available; the integration format."""

    format = "json"
    media_type = "application/json"
    extension = "json"

    def render(
        self, form: FormDefinition, session: FormSession, *, include_provenance: bool = True
    ) -> bytes:
        completeness = analyse(form, session)
        payload: dict[str, Any] = {
            "form": {
                "name": form.name,
                "version": form.version,
                "title": form.title,
            },
            "session": {
                "id": session.id,
                "status": session.status.value,
                "participants": session.participants,
                "created_at": session.created_at,
                "updated_at": session.updated_at,
            },
            "completeness": completeness.summary(),
            "sections": [],
            "agreements": [a.model_dump(mode="json") for a in session.agreements],
            "questions": [r.model_dump(mode="json") for r in session.support_requests],
            "action_items": [a.model_dump(mode="json") for a in session.action_items],
            "consistency_findings": [
                f.model_dump(mode="json") for f in session.consistency_findings
            ],
            "notes": recommended_gaps_as_notes(form, session),
        }

        for section in form.ordered_sections():
            entries = []
            for field in section.fields:
                answer = session.answers.get(field.id)
                entry: dict[str, Any] = {
                    "id": field.id,
                    "label": field.label,
                    "type": field.type.value,
                    "importance": field.importance.value,
                    "value": _REDACTED
                    if (field.sensitive and answer and answer.value)
                    else (answer.value if answer else None),
                    "state": answer.state.value if answer else AnswerState.EMPTY.value,
                }
                if include_provenance and answer and answer.provenance:
                    entry["provenance"] = answer.provenance.model_dump(mode="json")
                    entry["confidence"] = answer.confidence
                entries.append(entry)
            payload["sections"].append(
                {"id": section.id, "title": section.title, "fields": entries}
            )

        return json.dumps(payload, indent=2, ensure_ascii=False, default=str).encode("utf-8")


class MarkdownRenderer(Renderer):
    """Readable, diffable, and always available. The review format."""

    format = "markdown"
    media_type = "text/markdown"
    extension = "md"

    def render(
        self, form: FormDefinition, session: FormSession, *, include_provenance: bool = True
    ) -> bytes:
        lines: list[str] = [
            f"# {form.title}",
            "",
            f"**Form:** `{form.name}` v{form.version}  ",
            f"**Session:** `{session.id}`  ",
            f"**Status:** {session.status.value.replace('_', ' ').title()}  ",
            f"**Contributors:** {', '.join(session.participants) or '—'}  ",
            f"**Completeness:** {completeness_line(form, session)}",
            "",
        ]
        if form.description:
            lines += [f"> {form.description}", ""]

        for section in form.ordered_sections():
            rows = [
                (field, session.answers.get(field.id))
                for field in section.fields
                if session.answers.get(field.id) is not None
                and session.answers[field.id].state is not AnswerState.EMPTY
            ]
            if not rows:
                continue
            lines += [f"## {section.title}", ""]
            if section.description:
                lines += [section.description, ""]
            lines += ["| Field | Value | Status |", "| --- | --- | --- |"]
            for field, answer in rows:
                value = _display(field, answer).replace("|", "\\|").replace("\n", "<br>")
                lines.append(f"| {field.label} | {value} | {_status_note(answer)} |")
            lines.append("")

        # Before the action items: what was agreed is the basis everything below
        # it rests on, and an approver reading a refusal wants to see it early
        # rather than in an appendix.
        agreed = agreement_rows(form, session)
        if agreed:
            lines += ["## Agreements", ""]
            lines += [
                "| Agreement | Type | Version | Decision | By | When |",
                "| --- | --- | --- | --- | --- | --- |",
            ]
            for title, kind, version, decision, actor, when in agreed:
                mark = "Accepted" if decision == "accepted" else "**Declined**"
                lines.append(f"| {title} | {kind} | {version} | {mark} | {actor} | {when} |")
            lines.append("")
            for record in session.agreements:
                agreement = form.try_agreement(record.agreement_id)
                heading = agreement.title if agreement else record.agreement_id
                # The words, not a summary of them. A record of consent that
                # does not carry what was consented to is a record of nothing.
                lines += [f"> **{heading}** — {record.text.strip()}", ""]

        items = open_items(session)
        if items:
            lines += ["## Open action items", ""]
            lines += ["| # | Action | Owner | Due |", "| --- | --- | --- | --- |"]
            for index, item in enumerate(items, start=1):
                lines.append(
                    f"| {index} | {item.description} | {item.owner or '—'} | {item.due_date or '—'} |"
                )
            lines.append("")

        # The reasons matter more than the discrepancies. "Low technical risk,
        # high business impact" is a perfectly sound pair, and the owner's
        # explanation of why is the thing an approver would otherwise have to
        # ask for in a separate round trip.
        noted = noted_discrepancies(session)
        if noted:
            lines += ["## Noted discrepancies", ""]
            lines += ["| What was queried | The owner's answer |", "| --- | --- |"]
            for finding in noted:
                message = finding.message.replace("|", "\\|")
                resolution = (finding.resolution or "—").replace("|", "\\|").replace("\n", "<br>")
                lines.append(f"| {message} | {resolution} |")
            lines.append("")

        unresolved = [f for f in session.consistency_findings if f.is_outstanding]
        if unresolved:
            lines += ["## Unresolved discrepancies", ""]
            lines += [
                f"- {f.message}" + (f" _({f.evidence})_" if f.evidence else "") for f in unresolved
            ]
            lines.append("")

        # A question the participant asked and nobody could answer is context an
        # approver needs: it says which values were settled on solid ground and
        # which were the best anyone could do while waiting to hear back.
        questions = open_questions(session)
        if questions:
            lines += ["## Open questions", ""]
            lines += ["| Question | About | With | Raised by |", "| --- | --- | --- | --- |"]
            for request in questions:
                labels = ", ".join(form.field(f).label for f in request.fields if form.try_field(f))
                text = request.question.replace("|", "\\|").replace("\n", " ")
                lines.append(
                    f"| {text} | {labels or '—'} | {request.team or '—'} | {request.asked_by} |"
                )
            lines.append("")

        notes = recommended_gaps_as_notes(form, session)
        if notes:
            lines += ["## Notes", ""] + [f"- {n}" for n in notes] + [""]

        if include_provenance:
            lines += self._provenance_lines(form, session)

        return "\n".join(lines).encode("utf-8")

    @staticmethod
    def _provenance_lines(form: FormDefinition, session: FormSession) -> list[str]:
        rows: list[str] = []
        for field in form.fields():
            answer = session.answers.get(field.id)
            if answer is None or answer.provenance is None or field.sensitive:
                continue
            p = answer.provenance
            # For a value the wording pass rewrote, the typed text is the
            # evidence that matters: it is what the person actually said, and
            # the polished version is right there in the table above.
            source = answer.raw_value if answer.polished else p.evidence
            evidence = (source or "").replace("|", "\\|").replace("\n", " ")[:180]
            rows.append(
                f"| {field.label} | {p.channel.value} | {p.author} | "
                f"{answer.confidence:.2f} | {evidence} |"
            )
        if not rows:
            return []
        return [
            "## Provenance",
            "",
            "Where each captured value came from.",
            "",
            "| Field | Source | Stated by | Confidence | Evidence |",
            "| --- | --- | --- | --- | --- |",
            *rows,
            "",
        ]


class ExcelRenderer(Renderer):
    """Excel workbook: a sheet per topic, plus actions and provenance."""

    format = "xlsx"
    media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    extension = "xlsx"

    @classmethod
    def available(cls) -> bool:
        try:
            import openpyxl  # noqa: F401

            return True
        except ImportError:
            return False

    def render(
        self, form: FormDefinition, session: FormSession, *, include_provenance: bool = True
    ) -> bytes:
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Alignment, Font, PatternFill
            from openpyxl.utils import get_column_letter
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise ConfigurationError(
                "xlsx rendering requires openpyxl (install sa-forms[xlsx])", cause=exc
            ) from exc

        workbook = Workbook()
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill("solid", fgColor="1F4E78")
        label_font = Font(bold=True)
        wrap = Alignment(wrap_text=True, vertical="top")

        def write_header(sheet: Any, headings: list[str]) -> None:
            for column, heading in enumerate(headings, start=1):
                cell = sheet.cell(row=1, column=column, value=heading)
                cell.font = header_font
                cell.fill = header_fill
            sheet.freeze_panes = "A2"

        # -- summary ------------------------------------------------------
        summary = workbook.active
        summary.title = "Summary"
        completeness = analyse(form, session)
        for row, (key, value) in enumerate(
            [
                ("Form", form.title),
                ("Form ID", f"{form.name} v{form.version}"),
                ("Session", session.id),
                ("Status", session.status.value.replace("_", " ").title()),
                ("Contributors", ", ".join(session.participants) or "—"),
                (
                    "Required captured",
                    f"{completeness.mandatory_answered}/{completeness.mandatory_total}",
                ),
                (
                    "Agreements accepted",
                    f"{completeness.agreements_accepted}/{completeness.agreements_required}",
                ),
                ("Overall complete", f"{completeness.overall_percent:.0f}%"),
                ("Open actions", str(len(open_items(session)))),
            ],
            start=1,
        ):
            summary.cell(row=row, column=1, value=key).font = label_font
            summary.cell(row=row, column=2, value=value)
        summary.column_dimensions["A"].width = 22
        summary.column_dimensions["B"].width = 60

        # -- one sheet per section ----------------------------------------
        for section in form.ordered_sections():
            # Excel sheet names are capped at 31 chars and reject several
            # punctuation marks; sanitise rather than let openpyxl throw.
            title = "".join(c for c in section.title if c not in "[]:*?/\\")[:31] or section.id[:31]
            sheet = workbook.create_sheet(title)
            write_header(sheet, ["Field", "Value", "Status", "Required"])

            for row, field in enumerate(section.fields, start=2):
                answer = session.answers.get(field.id)
                sheet.cell(row=row, column=1, value=field.label).font = label_font
                cell = sheet.cell(row=row, column=2, value=_display(field, answer))
                cell.alignment = wrap
                sheet.cell(row=row, column=3, value=_status_note(answer))
                sheet.cell(row=row, column=4, value="Yes" if field.is_mandatory else "No")

            for column, width in zip("ABCD", (34, 62, 16, 10), strict=True):
                sheet.column_dimensions[column].width = width

        # -- agreements ----------------------------------------------------
        agreed = agreement_rows(form, session)
        if agreed:
            sheet = workbook.create_sheet("Agreements")
            write_header(
                sheet, ["Agreement", "Type", "Version", "Decision", "Accepted by", "When", "Text"]
            )
            for row, (record, columns) in enumerate(
                zip(session.agreements, agreed, strict=True), start=2
            ):
                for index, value in enumerate(columns, start=1):
                    sheet.cell(row=row, column=index, value=value)
                # The words themselves travel with the decision. A spreadsheet
                # row saying "accepted" with nothing to read is not evidence.
                sheet.cell(row=row, column=7, value=record.text).alignment = wrap
            for column, width in zip("ABCDEFG", (34, 14, 10, 12, 22, 20, 90), strict=True):
                sheet.column_dimensions[column].width = width

        # -- open questions -------------------------------------------------
        questions = open_questions(session)
        if questions:
            sheet = workbook.create_sheet("Open Questions")
            write_header(sheet, ["Question", "About", "With", "Contact", "Raised by"])
            for row, request in enumerate(questions, start=2):
                labels = ", ".join(form.field(f).label for f in request.fields if form.try_field(f))
                sheet.cell(row=row, column=1, value=request.question).alignment = wrap
                sheet.cell(row=row, column=2, value=labels)
                sheet.cell(row=row, column=3, value=request.team)
                sheet.cell(row=row, column=4, value=request.contact)
                sheet.cell(row=row, column=5, value=request.asked_by)
            for column, width in zip("ABCDE", (70, 28, 24, 30, 20), strict=True):
                sheet.column_dimensions[column].width = width

        # -- action items --------------------------------------------------
        items = open_items(session)
        if items:
            sheet = workbook.create_sheet("Action Items")
            write_header(sheet, ["#", "Action", "Owner", "Due", "Source field"])
            for row, item in enumerate(items, start=2):
                sheet.cell(row=row, column=1, value=row - 1)
                sheet.cell(row=row, column=2, value=item.description).alignment = wrap
                sheet.cell(row=row, column=3, value=item.owner or "")
                sheet.cell(row=row, column=4, value=item.due_date or "")
                sheet.cell(row=row, column=5, value=item.source_field_id or "")
            for column, width in zip("ABCDE", (6, 70, 20, 16, 24), strict=True):
                sheet.column_dimensions[column].width = width

        # -- provenance ----------------------------------------------------
        if include_provenance:
            sheet = workbook.create_sheet("Provenance")
            write_header(sheet, ["Field", "Source", "Stated by", "Confidence", "Evidence"])
            row = 2
            for field in form.fields():
                answer = session.answers.get(field.id)
                if answer is None or answer.provenance is None or field.sensitive:
                    continue
                p = answer.provenance
                sheet.cell(row=row, column=1, value=field.label)
                sheet.cell(row=row, column=2, value=p.channel.value)
                sheet.cell(row=row, column=3, value=p.author)
                sheet.cell(row=row, column=4, value=round(answer.confidence, 2))
                sheet.cell(row=row, column=5, value=(p.evidence or "")[:500]).alignment = wrap
                row += 1
            for column, width in zip("ABCDE", (30, 14, 22, 12, 80), strict=True):
                sheet.column_dimensions[get_column_letter("ABCDE".index(column) + 1)].width = width

        buffer = io.BytesIO()
        workbook.save(buffer)
        return buffer.getvalue()


class PdfRenderer(Renderer):
    """PDF for distribution and sign-off."""

    format = "pdf"
    media_type = "application/pdf"
    extension = "pdf"

    @classmethod
    def available(cls) -> bool:
        try:
            import fpdf  # noqa: F401

            return True
        except ImportError:
            return False

    def render(
        self, form: FormDefinition, session: FormSession, *, include_provenance: bool = True
    ) -> bytes:
        try:
            from fpdf import FPDF
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise ConfigurationError(
                "pdf rendering requires fpdf2 (install sa-forms[pdf])", cause=exc
            ) from exc

        pdf = FPDF(orientation="P", unit="mm", format="A4")
        pdf.set_auto_page_break(auto=True, margin=16)
        pdf.add_page()

        def text(value: str) -> str:  # - the table is *about* these characters
            # The core PDF fonts are Latin-1 only; substitute rather than raise
            # on a smart quote or an em dash pasted in from email.
            return (
                value.replace("—", "-")
                .replace("–", "-")
                .replace("’", "'")
                .replace("‘", "'")
                .replace("“", '"')
                .replace("”", '"')
                .replace("•", "-")
                .replace("…", "...")
                .encode("latin-1", "replace")
                .decode("latin-1")
            )

        # Widths are computed from the page and passed explicitly. Relying on
        # `multi_cell(w=0)` ("extend to the right margin") is what broke here:
        # after a two-column row the cursor sits near the right edge, so the
        # next full-width cell resolves to a few millimetres and fpdf raises
        # "not enough horizontal space". Always reset x, always pass a width.
        usable = pdf.w - pdf.l_margin - pdf.r_margin
        label_width = usable * 0.34
        value_width = usable - label_width

        def full_width(content: str, height: float = 5.0) -> None:
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(usable, height, text(content))

        def heading(label: str, size: int = 13) -> None:
            pdf.set_font("Helvetica", "B", size)
            pdf.set_text_color(31, 78, 120)
            full_width(label, 7)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(1)

        # -- title block ---------------------------------------------------
        pdf.set_font("Helvetica", "B", 18)
        full_width(form.title, 10)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(90, 90, 90)
        full_width(
            f"{form.name} v{form.version}  |  Session {session.id}  |  "
            f"{session.status.value.replace('_', ' ').title()}\n"
            f"Contributors: {', '.join(session.participants) or '-'}  |  "
            f"{completeness_line(form, session)}"
        )
        pdf.set_text_color(0, 0, 0)
        pdf.ln(4)

        # -- answers -------------------------------------------------------
        for section in form.ordered_sections():
            rows = [
                (f, session.answers.get(f.id))
                for f in section.fields
                if session.answers.get(f.id) is not None
                and session.answers[f.id].state is not AnswerState.EMPTY
            ]
            if not rows:
                continue
            heading(section.title)
            for field, answer in rows:
                start_y = pdf.get_y()
                pdf.set_xy(pdf.l_margin, start_y)
                pdf.set_font("Helvetica", "B", 9)
                pdf.multi_cell(label_width, 5, text(field.label))
                label_end = pdf.get_y()

                pdf.set_xy(pdf.l_margin + label_width, start_y)
                pdf.set_font("Helvetica", "", 9)
                pdf.multi_cell(value_width, 5, text(_display(field, answer) or "-"))
                # Re-align both columns and return x to the margin, so the next
                # row (or section heading) starts from a known position.
                pdf.set_y(max(label_end, pdf.get_y()))
                pdf.set_x(pdf.l_margin)
            pdf.ln(3)

        # -- agreements ----------------------------------------------------
        if session.agreements:
            heading("Agreements")
            for record in session.agreements:
                agreement = form.try_agreement(record.agreement_id)
                pdf.set_font("Helvetica", "B", 9)
                verdict = "Accepted" if record.accepted else "DECLINED"
                full_width(
                    f"{agreement.title if agreement else record.agreement_id} "
                    f"(v{record.version}) - {verdict} by {record.actor}, "
                    f"{_stamp(record.decided_at)}",
                    4.5,
                )
                pdf.set_font("Helvetica", "", 8)
                pdf.set_text_color(90, 90, 90)
                full_width(record.text, 4.5)
                pdf.set_text_color(0, 0, 0)
                pdf.ln(1)
            pdf.ln(2)

        # -- open questions -------------------------------------------------
        questions = open_questions(session)
        if questions:
            heading("Open questions")
            pdf.set_font("Helvetica", "", 9)
            for request in questions:
                labels = ", ".join(form.field(f).label for f in request.fields if form.try_field(f))
                about = f" (about {labels})" if labels else ""
                full_width(f"- \"{request.question}\"{about} - with {request.team or 'the owner'}")
            pdf.ln(3)

        # -- actions -------------------------------------------------------
        items = open_items(session)
        if items:
            heading("Open action items")
            pdf.set_font("Helvetica", "", 9)
            for index, item in enumerate(items, start=1):
                suffix = []
                if item.owner:
                    suffix.append(f"owner: {item.owner}")
                if item.due_date:
                    suffix.append(f"due: {item.due_date}")
                tail = f" ({'; '.join(suffix)})" if suffix else ""
                full_width(f"{index}. {item.description}{tail}")
            pdf.ln(3)

        # -- provenance ----------------------------------------------------
        if include_provenance:
            entries = [
                (f, session.answers[f.id])
                for f in form.fields()
                if session.answers.get(f.id) is not None
                and session.answers[f.id].provenance is not None
                and not f.sensitive
            ]
            if entries:
                pdf.add_page()
                heading("Provenance")
                pdf.set_font("Helvetica", "", 8)
                full_width("Where each captured value came from.")
                pdf.ln(2)
                for field, answer in entries:
                    p = answer.provenance
                    assert p is not None  # filtered above
                    pdf.set_font("Helvetica", "B", 8)
                    full_width(field.label, 4.5)
                    pdf.set_font("Helvetica", "", 8)
                    pdf.set_text_color(90, 90, 90)
                    full_width(
                        f"{p.channel.value} - {p.author} - confidence "
                        f"{answer.confidence:.2f}\n\"{(p.evidence or '')[:300]}\"",
                        4.5,
                    )
                    pdf.set_text_color(0, 0, 0)
                    pdf.ln(1)

        output = pdf.output()
        return bytes(output) if not isinstance(output, bytes) else output


class DocxRenderer(Renderer):
    """Word document, for organisations whose sign-off runs on Word."""

    format = "docx"
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    extension = "docx"

    @classmethod
    def available(cls) -> bool:
        try:
            import docx  # noqa: F401

            return True
        except ImportError:
            return False

    def render(
        self, form: FormDefinition, session: FormSession, *, include_provenance: bool = True
    ) -> bytes:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise ConfigurationError(
                "docx rendering requires python-docx (install sa-forms[docx])", cause=exc
            ) from exc

        document = Document()
        document.add_heading(form.title, level=0)
        document.add_paragraph(
            f"{form.name} v{form.version} | Session {session.id} | "
            f"{session.status.value.replace('_', ' ').title()}"
        )
        document.add_paragraph(
            f"Contributors: {', '.join(session.participants) or '—'} | "
            f"{completeness_line(form, session)}"
        )

        for section in form.ordered_sections():
            rows = [
                (f, session.answers.get(f.id))
                for f in section.fields
                if session.answers.get(f.id) is not None
                and session.answers[f.id].state is not AnswerState.EMPTY
            ]
            if not rows:
                continue
            document.add_heading(section.title, level=1)
            table = document.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            headings = table.rows[0].cells
            headings[0].text, headings[1].text, headings[2].text = "Field", "Value", "Status"
            for field, answer in rows:
                cells = table.add_row().cells
                cells[0].text = field.label
                cells[1].text = _display(field, answer)
                cells[2].text = _status_note(answer)

        if session.agreements:
            document.add_heading("Agreements", level=1)
            table = document.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            headings = table.rows[0].cells
            for index, name in enumerate(("Agreement", "Decision", "By", "When")):
                headings[index].text = name
            for record in session.agreements:
                agreement = form.try_agreement(record.agreement_id)
                cells = table.add_row().cells
                cells[0].text = agreement.title if agreement else record.agreement_id
                cells[1].text = "Accepted" if record.accepted else "Declined"
                cells[2].text = record.actor
                cells[3].text = _stamp(record.decided_at)
            for record in session.agreements:
                document.add_paragraph(record.text, style="Intense Quote")

        questions = open_questions(session)
        if questions:
            document.add_heading("Open questions", level=1)
            for request in questions:
                document.add_paragraph(
                    f"“{request.question}” — with {request.team or 'the form owner'}",
                    style="List Bullet",
                )

        items = open_items(session)
        if items:
            document.add_heading("Open action items", level=1)
            for item in items:
                suffix = f" (owner: {item.owner})" if item.owner else ""
                document.add_paragraph(f"{item.description}{suffix}", style="List Number")

        if include_provenance:
            entries = [
                (f, session.answers[f.id])
                for f in form.fields()
                if session.answers.get(f.id) is not None
                and session.answers[f.id].provenance is not None
                and not f.sensitive
            ]
            if entries:
                document.add_heading("Provenance", level=1)
                table = document.add_table(rows=1, cols=4)
                table.style = "Light Grid Accent 1"
                headings = table.rows[0].cells
                for index, name in enumerate(("Field", "Source", "Stated by", "Evidence")):
                    headings[index].text = name
                for field, answer in entries:
                    p = answer.provenance
                    assert p is not None  # filtered above
                    cells = table.add_row().cells
                    cells[0].text = field.label
                    cells[1].text = p.channel.value
                    cells[2].text = p.author
                    cells[3].text = (p.evidence or "")[:300]

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()


_RENDERERS: dict[str, type[Renderer]] = {
    "json": JsonRenderer,
    "markdown": MarkdownRenderer,
    "md": MarkdownRenderer,
    "xlsx": ExcelRenderer,
    "excel": ExcelRenderer,
    "pdf": PdfRenderer,
    "docx": DocxRenderer,
    "word": DocxRenderer,
}


def get_renderer(fmt: str) -> Renderer:
    factory = _RENDERERS.get(fmt.lower().lstrip("."))
    if factory is None:
        raise ConfigurationError(
            f"unsupported artifact format '{fmt}'",
            details={"format": fmt, "supported": available_formats()},
        )
    if not factory.available():
        raise ConfigurationError(
            f"the '{fmt}' renderer is not installed on this deployment",
            details={"format": fmt, "available": available_formats()},
        )
    return factory()


def available_formats() -> list[str]:
    """Formats this deployment can actually produce."""
    seen: dict[type[Renderer], str] = {}
    for name, factory in _RENDERERS.items():
        if factory.available() and factory not in seen:
            seen[factory] = name
    return sorted(seen.values())


def render_session(
    form: FormDefinition,
    session: FormSession,
    fmt: str,
    *,
    include_provenance: bool = True,
) -> tuple[bytes, str, str]:
    """Render a session. Returns ``(content, filename, media_type)``."""
    renderer = get_renderer(fmt)
    content = renderer.render(form, session, include_provenance=include_provenance)
    filename = f"{form.name}_{session.id}.{renderer.extension}"
    logger.info(
        "rendered artifact",
        extra={
            "form": form.name,
            "session": session.id,
            "format": renderer.format,
            "bytes": len(content),
        },
    )
    return content, filename, renderer.media_type


def artifact_status_for(session: FormSession) -> ArtifactStatus:
    """Map session status onto the artifact's initial status."""
    from .models import SessionStatus

    return {
        SessionStatus.APPROVED: ArtifactStatus.APPROVED,
        SessionStatus.BASELINED: ArtifactStatus.BASELINED,
        SessionStatus.IN_REVIEW: ArtifactStatus.IN_REVIEW,
    }.get(session.status, ArtifactStatus.DRAFT)


__all__ = [
    "DocxRenderer",
    "ExcelRenderer",
    "JsonRenderer",
    "MarkdownRenderer",
    "PdfRenderer",
    "Renderer",
    "artifact_status_for",
    "available_formats",
    "get_renderer",
    "render_session",
]
