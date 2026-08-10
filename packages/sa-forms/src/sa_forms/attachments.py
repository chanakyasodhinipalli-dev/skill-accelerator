"""Email and attachment parsing.

The premise of this package is that the information needed to fill a form
already exists in a conversation. Very often it exists in the *attachment* to
that conversation — the spreadsheet of impacted systems, the runbook, the design
note — and asking someone to retype it is exactly the manual work this is meant
to remove.

So an ``.eml`` file with three attachments becomes one email message plus one
message per attachment, all normalised to :class:`SourceMessage` and run through
the same extraction pipeline as anything else. Nothing downstream knows the
difference.

Extraction is per format and best-effort. A file that cannot be read as text —
an image, a zip, a signed blob — is *recorded* rather than dropped: the artifact
should say "there was an attachment I could not read", because silently ignoring
it is how a form ends up confidently wrong.
"""

from __future__ import annotations

import csv
import html
import io
import json
import re
from dataclasses import dataclass, field
from email import policy
from email.message import EmailMessage
from email.parser import BytesParser
from typing import Any

from sa_platform.errors import ValidationError
from sa_platform.logging import get_logger

logger = get_logger(__name__)

#: Attachments above this are not parsed. The cost is unbounded memory and a
#: very long extraction pass, for a file that is almost never form content.
MAX_ATTACHMENT_BYTES = 25 * 1024 * 1024
#: Text kept per attachment. A 400-page PDF contributes its opening, not itself.
MAX_TEXT_CHARS = 60_000
#: Rows read from a spreadsheet or CSV.
MAX_TABLE_ROWS = 400

_HTML_BLOCK = re.compile(r"(?i)</(p|div|tr|li|h[1-6]|table|section)>|<br\s*/?>")
_HTML_DROP = re.compile(r"(?is)<(script|style|head)[^>]*>.*?</\1>")
_HTML_TAG = re.compile(r"<[^>]+>")

#: Content types that are never form content, so they are recorded and skipped
#: rather than run through a decoder that will produce mojibake.
_OPAQUE_PREFIXES = ("image/", "audio/", "video/", "application/zip", "application/x-")


@dataclass(slots=True)
class Attachment:
    """One attached file and whatever text could be read out of it."""

    filename: str
    content_type: str = "application/octet-stream"
    size_bytes: int = 0
    text: str = ""
    #: Why there is no text, when there is none. Surfaced to the user.
    note: str = ""

    @property
    def extracted(self) -> bool:
        return bool(self.text.strip())

    def to_dict(self) -> dict[str, Any]:
        return {
            "filename": self.filename,
            "content_type": self.content_type,
            "size_bytes": self.size_bytes,
            "extracted": self.extracted,
            "characters": len(self.text),
            "note": self.note,
        }


@dataclass(slots=True)
class ParsedEmail:
    """An email reduced to the parts that matter for extraction."""

    subject: str = ""
    sender: str = ""
    recipients: list[str] = field(default_factory=list)
    date: str = ""
    message_id: str | None = None
    body: str = ""
    attachments: list[Attachment] = field(default_factory=list)

    def summary(self) -> dict[str, Any]:
        return {
            "subject": self.subject,
            "from": self.sender,
            "recipients": self.recipients,
            "date": self.date,
            "body_characters": len(self.body),
            "attachments": [a.to_dict() for a in self.attachments],
        }


# ---------------------------------------------------------------------------
# Email
# ---------------------------------------------------------------------------


def parse_email(raw: bytes | str, *, extract_attachments: bool = True) -> ParsedEmail:
    """Parse RFC 5322 bytes (an ``.eml`` file) into its parts."""
    data = raw.encode("utf-8", "replace") if isinstance(raw, str) else raw
    if _looks_like_outlook_msg(data):
        raise ValidationError(
            "this looks like an Outlook .msg file, which is a compound binary format. "
            "Save or export the mail as .eml (Save As -> Outlook Message Format is .msg; "
            "use 'Save as type: Text/MIME'), or forward it to a mailbox that stores MIME.",
            details={"format": "application/vnd.ms-outlook"},
        )

    try:
        message: EmailMessage = BytesParser(policy=policy.default).parsebytes(data)
    except Exception as exc:  # noqa: BLE001 - any malformed MIME lands here
        raise ValidationError(f"could not parse the email: {exc}", cause=exc) from exc

    parsed = ParsedEmail(
        subject=str(message.get("Subject", "") or ""),
        sender=str(message.get("From", "") or ""),
        recipients=[
            address.strip()
            for header in ("To", "Cc")
            for address in str(message.get(header, "") or "").split(",")
            if address.strip()
        ],
        date=str(message.get("Date", "") or ""),
        message_id=str(message.get("Message-ID", "") or "") or None,
        body=_body_text(message),
    )

    if extract_attachments:
        for part in message.iter_attachments():  # type: ignore[union-attr]
            attachment = _attachment_from_part(part)
            if attachment is not None:
                parsed.attachments.append(attachment)

    logger.info(
        "parsed email",
        extra={
            "subject": parsed.subject[:120],
            "attachments": len(parsed.attachments),
            "body_characters": len(parsed.body),
        },
    )
    return parsed


def _looks_like_outlook_msg(data: bytes) -> bool:
    # OLE2 compound-file magic. Detected explicitly so the user gets an
    # actionable message instead of a wall of binary treated as a body.
    return data[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def _body_text(message: EmailMessage) -> str:
    """Prefer the plain-text alternative; fall back to converting the HTML one."""
    try:
        plain = message.get_body(preferencelist=("plain",))
        if plain is not None:
            return _decode_part(plain)
        rich = message.get_body(preferencelist=("html",))
        if rich is not None:
            return html_to_text(_decode_part(rich))
    except Exception as exc:  # noqa: BLE001 - malformed multipart
        logger.warning("could not read the email body", extra={"error": str(exc)})
    return ""


def _decode_part(part: Any) -> str:
    payload = part.get_payload(decode=True)
    if payload is None:
        content = part.get_payload()
        return content if isinstance(content, str) else ""
    charset = part.get_content_charset() or "utf-8"
    return payload.decode(charset, errors="replace")


def _attachment_from_part(part: Any) -> Attachment | None:
    filename = part.get_filename() or "attachment"
    content_type = part.get_content_type()
    payload = part.get_payload(decode=True)
    if payload is None:
        return None
    return extract_attachment(filename, payload, content_type=content_type)


# ---------------------------------------------------------------------------
# Attachment text extraction
# ---------------------------------------------------------------------------


def extract_attachment(filename: str, content: bytes, *, content_type: str = "") -> Attachment:
    """Read whatever text an attachment holds. Never raises."""
    attachment = Attachment(
        filename=filename,
        content_type=content_type or _guess_content_type(filename),
        size_bytes=len(content),
    )

    if len(content) > MAX_ATTACHMENT_BYTES:
        attachment.note = (
            f"skipped: {len(content) // 1024 // 1024} MB exceeds the "
            f"{MAX_ATTACHMENT_BYTES // 1024 // 1024} MB extraction limit"
        )
        return attachment

    if attachment.content_type.startswith(_OPAQUE_PREFIXES):
        attachment.note = f"not text-extractable ({attachment.content_type})"
        return attachment

    suffix = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    reader = _READERS.get(suffix)
    if reader is None:
        reader = _read_plain if _looks_textual(content) else None
    if reader is None:
        attachment.note = f"no reader for '.{suffix}'" if suffix else "unrecognised format"
        return attachment

    try:
        text = reader(content)
    except _MissingDependencyError as exc:
        attachment.note = str(exc)
        return attachment
    except Exception as exc:  # noqa: BLE001 - one bad file must not fail the ingest
        logger.warning(
            "attachment extraction failed",
            # Not "filename": `logging` reserves it on a LogRecord.
            extra={"attachment": filename, "error": str(exc)},
        )
        attachment.note = f"could not be read: {exc}"
        return attachment

    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS]
        attachment.note = f"truncated to {MAX_TEXT_CHARS} characters"
    attachment.text = text.strip()
    if not attachment.text and not attachment.note:
        attachment.note = "no text content"
    return attachment


class _MissingDependencyError(RuntimeError):
    """An optional reader library is not installed."""


def _guess_content_type(filename: str) -> str:
    import mimetypes

    return mimetypes.guess_type(filename)[0] or "application/octet-stream"


def _looks_textual(content: bytes) -> bool:
    """Decide whether an unknown extension is worth decoding as text."""
    sample = content[:2048]
    if b"\x00" in sample:
        return False
    printable = sum(1 for byte in sample if 32 <= byte < 127 or byte in (9, 10, 13))
    return bool(sample) and printable / len(sample) > 0.85


def _read_plain(content: bytes) -> str:
    return content.decode("utf-8", errors="replace")


def _read_html(content: bytes) -> str:
    return html_to_text(_read_plain(content))


def _read_json(content: bytes) -> str:
    data = json.loads(_read_plain(content))
    return json.dumps(data, indent=2, default=str)


def _read_csv(content: bytes) -> str:
    """Render a CSV as labelled lines.

    ``Label: value`` is the shape the deterministic extractor already matches,
    so a two-column key/value sheet is read without a model being involved.
    """
    text = _read_plain(content)
    try:
        dialect: Any = csv.Sniffer().sniff(text[:4096])
    except csv.Error:
        dialect = csv.excel
    rows = list(csv.reader(io.StringIO(text), dialect))
    return _render_rows(rows)


def _read_xlsx(content: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise _MissingDependencyError("install openpyxl to read .xlsx attachments") from exc

    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    blocks: list[str] = []
    try:
        for sheet in workbook.worksheets:
            rows = [
                ["" if cell is None else str(cell) for cell in row]
                for row in sheet.iter_rows(max_row=MAX_TABLE_ROWS, values_only=True)
            ]
            rendered = _render_rows(rows)
            if rendered:
                blocks.append(f"# Sheet: {sheet.title}\n{rendered}")
    finally:
        workbook.close()
    return "\n\n".join(blocks)


def _read_docx(content: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise _MissingDependencyError("install python-docx to read .docx attachments") from exc

    document = docx.Document(io.BytesIO(content))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise _MissingDependencyError("install pypdf to read .pdf attachments") from exc

    reader = PdfReader(io.BytesIO(content))
    return "\n\n".join((page.extract_text() or "") for page in reader.pages)


def _read_eml(content: bytes) -> str:
    """A forwarded mail carried as an attachment. One level only.

    Recursing further would let a mail chain of attached mail chains expand
    without bound; one level covers the real case (a forward) and stops there.
    """
    nested = parse_email(content, extract_attachments=False)
    header = f"Forwarded message from {nested.sender}: {nested.subject}"
    return f"{header}\n\n{nested.body}"


_READERS: dict[str, Any] = {
    "txt": _read_plain,
    "log": _read_plain,
    "md": _read_plain,
    "yaml": _read_plain,
    "yml": _read_plain,
    "json": _read_json,
    "csv": _read_csv,
    "tsv": _read_csv,
    "html": _read_html,
    "htm": _read_html,
    "xlsx": _read_xlsx,
    "xlsm": _read_xlsx,
    "docx": _read_docx,
    "pdf": _read_pdf,
    "eml": _read_eml,
}

#: Extensions with a reader, for the UI's file picker and the docs.
SUPPORTED_ATTACHMENT_TYPES = tuple(sorted(_READERS))


def _render_rows(rows: list[list[str]]) -> str:
    """Render tabular data so the extractor can read it.

    Two columns become ``Label: value`` lines, which the deterministic matcher
    already understands. Wider tables keep a header row and pipe-separate, which
    reads as a table to a model without inventing a Markdown renderer.
    """
    rows = [r for r in rows[:MAX_TABLE_ROWS] if any(str(c).strip() for c in r)]
    if not rows:
        return ""

    width = max(len(r) for r in rows)
    if width <= 2:
        return "\n".join(
            f"{str(r[0]).strip()}: {str(r[1]).strip()}" if len(r) > 1 else str(r[0]).strip()
            for r in rows
        )
    return "\n".join(" | ".join(str(c).strip() for c in row) for row in rows)


def html_to_text(markup: str) -> str:
    """Convert HTML to readable text.

    Deliberately small: this handles mail bodies, which are structurally simple,
    and adding an HTML parser dependency to read a signature block is not a
    trade worth making.
    """
    text = _HTML_DROP.sub(" ", markup)
    text = _HTML_BLOCK.sub("\n", text)
    text = _HTML_TAG.sub(" ", text)
    text = html.unescape(text)
    # The tab and non-breaking space are written as escapes: a literal NBSP
    # in source is invisible, and mail bodies are full of them.
    text = re.sub(r"[ \t\xa0]+", " ", text)
    # An opening tag collapses to a space, which would otherwise indent every
    # line that follows one — and an indented `Label: value` stops matching.
    text = "\n".join(line.strip() for line in text.splitlines())
    return re.sub(r"\n\s*\n\s*\n+", "\n\n", text).strip()


__all__ = [
    "MAX_ATTACHMENT_BYTES",
    "MAX_TABLE_ROWS",
    "MAX_TEXT_CHARS",
    "SUPPORTED_ATTACHMENT_TYPES",
    "Attachment",
    "ParsedEmail",
    "extract_attachment",
    "html_to_text",
    "parse_email",
]
